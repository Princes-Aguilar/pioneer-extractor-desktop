import json
import sys
from docx import Document

def parse_number(v):
    if v is None:
        return 0.0
    s = str(v).replace(",", "").strip()
    out = ""
    dot = False
    for ch in s:
        if ch.isdigit():
            out += ch
        elif ch == "." and not dot:
            out += "."
            dot = True
        elif out:
            break
    try:
        return float(out) if out else 0.0
    except:
        return 0.0

def fmt2(x):
    try:
        return f"{float(x or 0):,.2f}"
    except:
        return "0.00"

def replace_all(doc, mapping):
    for p in doc.paragraphs:
        for key, val in mapping.items():
            if key in p.text:
                p.text = p.text.replace(key, val)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, val in mapping.items():
                        if key in p.text:
                            p.text = p.text.replace(key, val)

def main():
    if len(sys.argv) != 4:
        print("Usage: script template out payload", file=sys.stderr)
        sys.exit(2)

    template_path = sys.argv[1]
    out_path = sys.argv[2]
    payload_path = sys.argv[3]

    with open(payload_path, "r", encoding="utf-8") as f:
        payload = json.load(f)

    items = payload.get("items") or []

    dg_items = payload.get("dgItems") or []
    ndg_items = payload.get("ndgItems") or []

    if not dg_items and not ndg_items:
        for it in items:
            name = (
                it.get("properShippingName")
                or it.get("description")
                or it.get("productName")
                or "—"
            )
            dg = str(it.get("dgStatus","")).strip().upper()
            if dg in ("DG","YES","Y","TRUE"):
                dg_items.append(name)
            else:
                ndg_items.append(name)

    total_gw = payload.get("totalGrossWeightKgs")
    if total_gw is None:
        total_gw = 0.0
        for it in items:
            total_gw += parse_number(it.get("grossWeight"))

    dg_text = "\n".join([f"{i+1}. {x}" for i, x in enumerate(dg_items)]) or "—"
    ndg_text = "\n".join([f"{i+1}. {x}" for i, x in enumerate(ndg_items)]) or "—"

    mapping = {
        "{{DG_LIST}}": dg_text,
        "{{NDG_LIST}}": ndg_text,
        "{{TOTAL_GW}}": fmt2(total_gw),
    }

    doc = Document(template_path)
    replace_all(doc, mapping)
    doc.save(out_path)
    print(out_path)

if __name__ == "__main__":
    main()