import json
import sys
from docx import Document

def upper(s):
    return str(s or "").strip().upper()

def parse_number(v):
    if v is None:
        return 0.0
    s = str(v).replace(",", "").strip()
    out = ""
    dot = False
    for ch in s:
        if ch.isdigit():
            out += ch
        elif ch == "." and not dot:
            out += "."
            dot = True
        elif out:
            break
    try:
        return float(out) if out else 0.0
    except:
        return 0.0

def _replace_in_paragraph_runs(paragraph, mapping):
    if not paragraph.runs:
        return
    full = "".join(run.text for run in paragraph.runs)
    new = full
    for k, v in mapping.items():
        if k in new:
            new = new.replace(k, v)
    if new != full:
        paragraph.runs[0].text = new
        for r in paragraph.runs[1:]:
            r.text = ""

def replace_all(doc, mapping):
    for p in doc.paragraphs:
        _replace_in_paragraph_runs(p, mapping)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph_runs(p, mapping)

def main():
    if len(sys.argv) != 4:
        print("Usage: generate_nondg_cert_docx_cli.py <template.docx> <out.docx> <payload.json>", file=sys.stderr)
        sys.exit(2)

    template_path = sys.argv[1]
    out_path = sys.argv[2]
    payload_path = sys.argv[3]

    with open(payload_path, "r", encoding="utf-8") as f:
        payload = json.load(f)

    rows = payload.get("items") or []

    # Prefer explicit list if UI sends it
    ndg_items = payload.get("ndgItems") or []

    # Otherwise infer Non-DG from rows (DG status)
    if not ndg_items:
        for r in rows:
            status = upper(r.get("dgStatus"))
            name = (r.get("properShippingName") or r.get("description") or r.get("productName") or "—").strip()
            if status not in ("DG", "YES", "Y", "TRUE"):
                ndg_items.append(name)

    # Build list text
    ndg_list_text = "\n".join([f"{i+1}. {x}" for i, x in enumerate(ndg_items)]) or "—"

    # ✅ Supports either placeholder style OR your current NDG item#1 template
    mapping = {
        "{{NDG_LIST}}": ndg_list_text,      # if you choose to add this placeholder later
        "NDG item#1": ndg_list_text,        # replace first slot with full list
        "NDG item#2": "",                   # clear the unused slots
        "NDG item#3": "",
    }

    doc = Document(template_path)
    replace_all(doc, mapping)
    doc.save(out_path)
    print(out_path)

if __name__ == "__main__":
    main()