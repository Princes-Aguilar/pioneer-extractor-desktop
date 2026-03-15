import json
import sys
from datetime import datetime
from docx import Document

def parse_number(v):
    if v is None:
        return 0.0
    s = str(v).replace(",", "").strip()
    out = ""
    dot = False
    started = False

    for ch in s:
        if ch.isdigit():
            out += ch
            started = True
        elif ch == "." and not dot:
            out += "."
            dot = True
            started = True
        elif started:
            break

    try:
        return float(out) if out else 0.0
    except:
        return 0.0

def f2(x):
    try:
        return f"{float(x or 0):,.2f}"
    except:
        return "0.00"

def _replace_in_paragraph_runs(paragraph, mapping: dict):
    if not paragraph.runs:
        return

    full = "".join(run.text for run in paragraph.runs)
    new = full
    for k, v in mapping.items():
        if k in new:
            new = new.replace(k, v)

    if new != full:
        paragraph.runs[0].text = new
        for r in paragraph.runs[1:]:
            r.text = ""

def replace_text_in_doc(doc: Document, mapping: dict):
    for p in doc.paragraphs:
        _replace_in_paragraph_runs(p, mapping)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph_runs(p, mapping)

def main():
    if len(sys.argv) != 4:
        print("Usage: generate_preadvise_docx_cli.py <template.docx> <out.docx> <payload.json>", file=sys.stderr)
        sys.exit(2)

    template_path = sys.argv[1]
    out_path = sys.argv[2]
    payload_path = sys.argv[3]

    with open(payload_path, "r", encoding="utf-8") as f:
        payload = json.load(f)

    today = datetime.now().strftime("%B %d, %Y").upper()

    cargo_weight = parse_number(payload.get("cargoWeightKgs"))
    tare_weight = parse_number(payload.get("tareWeightKgs"))
    vgmd_weight = cargo_weight + tare_weight

    mapping = {
        "{{DATE_OF_DELIVERY}}": today,
        "{{FIRST_PORT}}": str(payload.get("firstPort", "")),
        "{{SECOND_PORT}}": str(payload.get("secondPort", "")),
        "{{BOOKING_NUMBER}}": str(payload.get("bookingNumber", "")),
        "{{VESSEL_VOYAGE}}": str(payload.get("vesselVoyage", "")),
        "{{CONTAINER_SIZE_TYPE}}": str(payload.get("containerSizeType", "")),
        "{{CONTAINER_NUMBERS}}": str(payload.get("containerNumbers", "")),
        "{{SEAL_NUMBERS}}": str(payload.get("sealNumbers", "")),
        "{{TARE_WEIGHT_KGS}}": f2(tare_weight),
        "{{CARGO_WEIGHT_KGS}}": f2(cargo_weight),
        "{{VGMD_KGS}}": f2(vgmd_weight),
        "{{TRUCKER}}": str(payload.get("trucker", "")),
        "{{PLATE_NUMBER}}": str(payload.get("plateNumber", "")),
        "{{UNNO_IMO_CLASS}}": str(payload.get("unnoImoClass", "")),
    }

    doc = Document(template_path)
    replace_text_in_doc(doc, mapping)
    doc.save(out_path)
    print(out_path)

if __name__ == "__main__":
    main()